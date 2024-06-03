from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, Markup, \
    send_from_directory
import pymysql
import bcrypt
from datetime import datetime
import markdown
from concurrent.futures import ThreadPoolExecutor
from docx import Document
import re
import sys
import os

# 确保工程根目录在系统路径中
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from long_essay_generator.AI_Agents import Agent_outline, Agent_heading, Agent_paragraph
from postgraduate_info.answer import main_func

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # 替换为你的实际密钥

zhipuai_API_KEY = 'a1c585c7b8106e12b92ceae99026a2cd.frA1fHlR0O4lyba4'

# 数据库配置
config = {
    'host': '110.41.49.124',
    'port': 3306,
    'user': 'root',
    'password': '@HWSJKmimashi111',
    'db': 'user_data',
    'charset': 'utf8mb4'
}


# 连接到数据库
def connect_db():
    return pymysql.connect(
        host=config['host'],
        user=config['user'],
        password=config['password'],
        db=config['db'],
        charset=config['charset'],
        cursorclass=pymysql.cursors.DictCursor
    )


# 获取用户历史记录
def get_history(user_id):
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            sql = "SELECT id, question, answer, timestamp FROM history WHERE user_id = %s ORDER BY timestamp DESC"
            cursor.execute(sql, (user_id,))
            result = cursor.fetchall()
            return result
    finally:
        connection.close()


# 添加历史记录
def add_history(user_id, question, answer):
    connection = connect_db()
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        with connection.cursor() as cursor:
            sql = "INSERT INTO history (user_id, question, answer, timestamp) VALUES (%s, %s, %s, %s)"
            cursor.execute(sql, (user_id, question, answer, timestamp))
            connection.commit()
            return cursor.lastrowid
    finally:
        connection.close()


# 用户注册或登录
def register_or_login(username, password):
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            sql = "SELECT user_id, password_hash FROM users WHERE username = %s"
            cursor.execute(sql, (username,))
            result = cursor.fetchone()
            if result:
                if bcrypt.checkpw(password.encode('utf-8'), result['password_hash'].encode('utf-8')):
                    return ('login', result['user_id'])
                else:
                    return ('error', '密码错误')
            else:
                hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
                sql = "INSERT INTO users (username, password_hash) VALUES (%s, %s)"
                cursor.execute(sql, (username, hashed_password))
                connection.commit()
                return ('register', cursor.lastrowid)
    finally:
        connection.close()


# 创建表
def create_tables():
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            create_users_table = """
            CREATE TABLE IF NOT EXISTS users (
                user_id INT AUTO_INCREMENT PRIMARY KEY,
                username VARCHAR(255) UNIQUE NOT NULL,
                password_hash CHAR(60) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
            """
            create_history_table = """
            CREATE TABLE IF NOT EXISTS history (
                id INT AUTO_INCREMENT PRIMARY KEY,
                user_id INT NOT NULL,
                question TEXT NOT NULL,
                answer TEXT NOT NULL,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE
            );
            """
            cursor.execute(create_users_table)
            cursor.execute(create_history_table)
            connection.commit()
    finally:
        connection.close()


create_tables()


# Markdown 转 HTML 方法
def md_to_html(md_content):
    exts = ['markdown.extensions.extra', 'markdown.extensions.codehilite', 'markdown.extensions.tables',
            'markdown.extensions.toc']
    html = markdown.markdown(md_content, extensions=exts)
    return Markup(html)


def get_complex_outline(user_input):
    outline = Agent_outline.get_complex_outline(user_input, zhipuai_API_KEY)
    pattern = r'(^\d+\.\S*[^\n]*)|(\d+\.\d+\S*[^\n]*)|(\d+\.\d+\.\d+\S*[^\n]*)'
    matches = re.findall(pattern, outline, re.MULTILINE)
    sections = [''.join(match) for match in matches if any(match)]
    return sections


def get_simple_outline(user_input):
    outline = Agent_outline.get_simple_outline(user_input, zhipuai_API_KEY)
    pattern = r'(^\d+\.\S*[^\n]*)|(\d+\.\d+\S*[^\n]*)|(\d+\.\d+\.\d+\S*[^\n]*)'
    matches = re.findall(pattern, outline, re.MULTILINE)
    sections = [''.join(match) for match in matches if any(match)]
    return sections


def get_heading(user_input):
    heading = Agent_heading.get_heading(user_input, zhipuai_API_KEY)
    return heading


def get_paragraph(outline, subheading):
    paragraph = Agent_paragraph.get_paragraph(outline, subheading, zhipuai_API_KEY)
    return paragraph


def is_third_level_heading(sections, index):
    if index + 1 < len(sections) and re.match(r'^\d+\.\d+\.\d+', sections[index + 1]):
        return True
    return False


def is_second_level_heading(sections, index):
    if index + 1 < len(sections) and re.match(r'^\d+\.\d+', sections[index + 1]):
        return True
    return False


def generate_long_paper(user_input):
    with ThreadPoolExecutor() as executor:
        outline_future = executor.submit(get_simple_outline, user_input)
        heading_future = executor.submit(get_heading, user_input)

        outline = outline_future.result()
        heading = heading_future.result()

    doc = Document()
    doc.add_heading(heading, level=0)

    for index, section in enumerate(outline):
        section = section.strip()
        if re.search(r'^\d+\.\s*[^\d\s]', section):
            if is_second_level_heading(outline, index):
                doc.add_heading(section, level=1)
            else:
                paragraph = get_paragraph(outline, section)
                doc.add_paragraph(paragraph)
        elif re.search(r'^\d+\.\d+\.\d+', section):
            doc.add_heading(section, level=3)
            paragraph = get_paragraph(outline, section)
            doc.add_paragraph(paragraph)
        elif re.search(r'^\d+\.\d+\s*[^\d\s]', section):
            if is_third_level_heading(outline, index):
                doc.add_heading(section, level=2)
            else:
                doc.add_heading(section, level=2)
                paragraph = get_paragraph(outline, section)
                doc.add_paragraph(paragraph)

    docx_path = f"output/{heading}.docx"
    doc.save(docx_path)
    return docx_path


def generate_longlong_paper(user_input):
    with ThreadPoolExecutor() as executor:
        outline_future = executor.submit(get_complex_outline, user_input)
        heading_future = executor.submit(get_heading, user_input)

        outline = outline_future.result()
        heading = heading_future.result()

    doc = Document()
    doc.add_heading(heading, level=0)

    for index, section in enumerate(outline):
        section = section.strip()
        if re.search(r'^\d+\.\s*[^\d\s]', section):
            if is_second_level_heading(outline, index):
                doc.add_heading(section, level=1)
            else:
                paragraph = get_paragraph(outline, section)
                doc.add_paragraph(paragraph)
        elif re.search(r'^\d+\.\d+\.\d+', section):
            doc.add_heading(section, level=3)
            paragraph = get_paragraph(outline, section)
            doc.add_paragraph(paragraph)
        elif re.search(r'^\d+\.\d+\s*[^\d\s]', section):
            if is_third_level_heading(outline, index):
                doc.add_heading(section, level=2)
            else:
                doc.add_heading(section, level=2)
                paragraph = get_paragraph(outline, section)
                doc.add_paragraph(paragraph)

    docx_path = f"output/{heading}.docx"
    doc.save(docx_path)
    return docx_path


@app.route('/')
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if not username:
            return jsonify({"status": "error", "message": "请输入您的账号"})
        if not password:
            return jsonify({"status": "error", "message": "请输入您的密码"})
        action, user_info = register_or_login(username, password)
        if action == 'login':
            session['serial_number'] = 0
            session['user_id'] = user_info
            session['username'] = username  # Add username to session
            return jsonify({"status": "success", "message": "登录成功"})
        elif action == 'register':
            session['user_id'] = user_info
            session['username'] = username  # Add username to session
            return jsonify({"status": "success", "message": "注册成功"})
        elif action == 'error':
            return jsonify({"status": "error", "message": user_info})

    return render_template('login.html')


@app.route('/home')
def home():
    if 'user_id' not in session:
        flash("请登录以查看主页。", "warning")
        return redirect(url_for('login'))

    user_history = get_history(session['user_id'])
    return render_template('home.html', history=user_history)


@app.route('/baoyan', methods=['POST'])
def baoyan():
    question = request.form.get('question')
    user_id = session.get('user_id')
    serial_number = session.get('serial_number')
    print('serial_number', serial_number)
    if not user_id:
        return jsonify({"error": "请登录以查看保研信息。"})

    response, answer_len = main_func(question, serial_number)

    session['serial_number'] += answer_len
    # 保存到历史记录表
    answer = "\n".join(response)
    add_history(user_id, question, answer)

    return jsonify({
        "results": response,
        "moreInfo": False
    })


@app.route('/baoyan')
def baoyan_page():
    if 'user_id' not in session:
        flash("请登录以查看保研信息。", "warning")
        return redirect(url_for('login'))

    user_history = get_history(session['user_id'])
    return render_template('baoyan.html', history=user_history)


@app.route('/shixi')
def shixi():
    if 'user_id' not in session:
        flash("请登录以查看实习信息。", "warning")
        return redirect(url_for('login'))

    user_history = get_history(session['user_id'])
    return render_template('shixi.html', history=user_history)


@app.route('/longtext')
def longtext():
    if 'user_id' not in session:
        flash("请登录以查看长文本生成页面。", "warning")
        return redirect(url_for('login'))

    user_history = get_history(session['user_id'])
    return render_template('longtext.html', history=user_history)


@app.route('/ask', methods=['POST'])
def ask():
    if 'user_id' not in session:
        return jsonify({"error": "请登录以保存您的问答记录。"})

    question = request.form.get('question')
    length = request.form.get('length')

    if question and length:
        if length == 'long':
            docx_path = generate_long_paper(question)
        elif length == 'longlong':
            docx_path = generate_longlong_paper(question)
        else:
            return jsonify({"error": "无效的长度选项。"})

        # 保存到历史记录表
        answer = f"生成的文档路径: {docx_path}"
        add_history(session['user_id'], question, answer)
        download_url = url_for('download', filename=os.path.basename(docx_path))
        return jsonify({"question": question, "url": download_url})

    return jsonify({"error": "未能获取问题或长度选项。"})


@app.route('/history/<int:record_id>')
def get_history_record(record_id):
    if 'user_id' not in session:
        return jsonify({"error": "请登录以查看历史记录。"})

    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            sql = "SELECT question, answer FROM history WHERE id = %s AND user_id = %s"
            cursor.execute(sql, (record_id, session['user_id']))
            result = cursor.fetchone()
            if result:
                result['answer'] = md_to_html(result['answer'])
                return jsonify(result)
            else:
                return jsonify({"error": "未找到历史记录。"})
    finally:
        connection.close()


@app.route('/download/<filename>', methods=['GET'])
def download(filename):
    return send_from_directory('output', filename)


@app.route('/wechat_login')
def wechat_login():
    return redirect(
        "https://open.weixin.qq.com/connect/qrconnect?appid=YOUR_APP_ID&redirect_uri=YOUR_REDIRECT_URI&response_type=code&scope=snsapi_login&state=STATE#wechat_redirect")


if __name__ == '__main__':
    if not os.path.exists('output'):
        os.makedirs('output')
    app.run(debug=True)
